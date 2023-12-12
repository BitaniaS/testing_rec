from sqlalchemy import create_engine, text,MetaData
from docx import Document
from sqlalchemy.orm import Session
from sqlalchemy.engine import reflection

# Connect to the database
engine = create_engine("mysql+pymysql://root:Asanti2002@localhost:3306/ipeds202021")
# Test the connection
# connection = engine.connect()

#Creating a metadata object and reflect the database schema. 
metadata = MetaData()
metadata.reflect(bind=engine)

# Loops to access information such as table name, columns, constraints and indexes for each table in metadata
#uncomment specific the outer for loop and then chosen inner for loop for spefici information from metadata

# for table in metadata.tables.values():
    # print(f"Table Name: {table.name}") 

    # for column in table.c:
    #     print(f"Column Name: {column.name}, Type: {column.type}")

    # for constraint in table.constraints:
    #     print(f"Table Name: {table.name}, Constraint: {constraint}")

    # for index in table.indexes:
    #     print(f"Table Name: {table.name}, Index: {index}")


# --------- BELOW IS CODE TO EXPORT TO WORD DOCUMENT (UNCOMMENT TO CHECK)------------------------------------------------ 
#path = 'C:/Users/bkete/OneDrive/Desktop/IS/testing_rec/my-rec-back-end/ipeds_data.docx' # change path to your preference
# #Create a new Word document
# doc = Document()
# # Create a session
# with Session(engine) as session:
#     result = session.execute(text("SELECT varName, longDescription FROM vartable20"))
#     for row in result:
#         line = f"variable {row[0]} description {row[1]},\n"
#         doc.add_paragraph(line)
#         # Process each row

# #save the word doc
# doc.save(path)
# # print("Data exported to Word document successfully.") // print statement once all data is exported to word

# ----------------END OF CODE TO EXPORT TO WORD -------------------------------------------------------------------------


# Create an inspector object
insp = reflection.Inspector.from_engine(engine) 
# Fetch names of tables and store in varaible table_names
table_names = insp.get_table_names()
# print("Table names:", table_names) #Print statement to check table_names

tables_not_used = ['filenames20','sectiontable20','tables20', 'valuesets20', 'vartable20']
reference_table = 'HD2020'
other_tables = {'F1920_F1A','F1920_F2','IC2020','IC2020_AY','DRVIC2020','SFA1920_P2','IC2020_PY','ADM2020','EF2020D','DRVF2020','DRVEF2020','IC2020MISSION'}
# all_tables = other_tables.add('HD2020')
other_tables.add(reference_table)
print(other_tables)
#function to look at duplicate and null UNITID values for tables 


conn = engine.connect()

 
'''
f19_a.F1C101, f19_f2.F2C07,f19_f2.F2C05,f19_f2.F2C06, f19_f2.F2C12, ic2.APPLFEEU,ic2.TUITVARY,ic2.TUITPL1,ic2.TUITPL2,ic2.TUITPL3,ic2.TUITPL4,
ic2_a.CHG3TGTD,ic2.STUSRV3,ic2.ROOM,ic2.BOARD, ic2.MEALSWK,ic2.ALLONCAM,ic2_a.chg1ay3,ic2_a.chg2ay3,ic2_a.chg3ay3,ic2_a.chg7py2,ic2.ROOMAMT,ic2.BOARDAMT,
ic2.RMBRDAMT,vic.CINDON,vic.CINSON,vic.COTSON,vic.CINDOFF,vic.CINSOFF,vic.COTSOFF,vic.CINDFAM,vic.CINSFAM,vic.COTSFAM,p2.NPGRN0,p2.sfNPT412,p2.NPT422,
p2.NPT432,p2.GRN4T30, p2.GRN4A30,p2.GRN4N40,ic2_p.CIPCODE1,ic2_p.CIPCODE2,ic2_p.CIPCODE3,ic2_p.CIPCODE4,ic2_p.CIPCODE5,ic2_p.CIPCODE6,hd.ADDR,hd.CITY,
hd.STABBR,hd.ZIP,hd.GROFFER, hd.CARNEGIE, ic2.OPENADMP,adm.ADMCON1,adm.ADMCON2,adm.ADMCON3,adm.ADMCON4,adm.ADMCON5,adm.ADMCON6,adm.ADMCON7,adm.ADMCON8,
adm.ADMCON9, adm.SATNUM,adm.SATPCT,adm.ACTNUM,adm.ACTPCT,adm.SATVR25, adm.SATVR75,adm.SATMT25,adm.SATMT75,adm.ACTCM25, adm.ACTCM75,adm.ACTEN25,adm.ACTEN75,
adm.ACTMT25, adm.ACTMT75,efd.RET_PCF,efd.STUFACR,vf.F3RSRCFT,ic2.ATHASSOC,ic2.ASSOC1,ic2.ASSOC2,ic2.SPORT1,ic2.CONFNO1,ic2.SPORT2,ic2.CONFNO2,ic2.SPORT3,
ic2.CONFNO3,ic2.SPORT4,ic2.CONFNO4,hd.LOCALE,hd.VETURL,ic2.VET1,ic2.VET2,ic2.VET3,ic2.VET4,ic2.VET5,ic2.VET9,efd.UGENTERN,vef.PctEnrWh,vef.PctEnrBK ,
vef.PctEnrHS,vef.PctEnrAP,vef.PctEnrAN,vef.PctEnrUn,vef.PctEnrNr,vef.PCTENRW,ic2.DISAB,ic2.DISABPCT,f19_a.F1C061,vf.F1STSVPC,vf.F1STSVFT,vf.F2STSVPC,
vf.F2STSVFT,hd.HBCU,hd.ACT,hd.CYACTIVE,hd.INSTNM,hd.CHFNM,hd.CHFTITLE, hd.GENTELE,hd.FAIDURL, hd.ADMINURL,hd.APPLURL,hd.DISAURL,hd.WEBADDR,hd.NPRICURL,
ic2.CNTLAFFI,ic2.RELAFFIL,ic2.STUSRV1, ic2.STUSRV2, ic2.STUSRV4,ic2.STUSRV8,ic2.STUSRV9, ic2_m.MISSIONURL, ic2.DISTCRS,ic2.DISTNCED, ic2.DSTNCED1,
ic2.DSTNCED3,ic2.DSTNUGC,ic2.DSTNUGP, ic2.DSTNUGN
'''

# SQL query
sql_query = text( """ 
CREATE TABLE new_combined_table AS
SELECT 
    f19_a.F1C101, f19_f2.F2C07,f19_f2.F2C05,f19_f2.F2C06, f19_f2.F2C12, ic2.APPLFEEU,ic2.TUITVARY,ic2.TUITPL1,ic2.TUITPL2,ic2.TUITPL3,ic2.TUITPL4,
    ic2_a.CHG3TGTD,ic2.STUSRV3,ic2.ROOM,ic2.BOARD, ic2.MEALSWK,ic2.ALLONCAM,ic2_a.chg1ay3,ic2_a.chg2ay3,ic2_a.chg3ay3,ic2_a.chg7py2,ic2.ROOMAMT,ic2.BOARDAMT,
    ic2.RMBRDAMT,vic.CINDON,vic.CINSON,vic.COTSON,vic.CINDOFF,vic.CINSOFF,vic.COTSOFF,vic.CINDFAM,vic.CINSFAM,vic.COTSFAM,p2.NPGRN0,p2.sfNPT412,p2.NPT422,
    p2.NPT432,p2.GRN4T30, p2.GRN4A30,p2.GRN4N40,ic2_p.CIPCODE1,ic2_p.CIPCODE2,ic2_p.CIPCODE3,ic2_p.CIPCODE4,ic2_p.CIPCODE5,ic2_p.CIPCODE6,hd.ADDR,hd.CITY,
    hd.STABBR,hd.ZIP,hd.GROFFER, hd.CARNEGIE, ic2.OPENADMP,adm.ADMCON1,adm.ADMCON2,adm.ADMCON3,adm.ADMCON4,adm.ADMCON5,adm.ADMCON6,adm.ADMCON7,adm.ADMCON8,
    adm.ADMCON9, adm.SATNUM,adm.SATPCT,adm.ACTNUM,adm.ACTPCT,adm.SATVR25, adm.SATVR75,adm.SATMT25,adm.SATMT75,adm.ACTCM25, adm.ACTCM75,adm.ACTEN25,adm.ACTEN75,
    adm.ACTMT25, adm.ACTMT75,efd.RET_PCF,efd.STUFACR,vf.F3RSRCFT,ic2.ATHASSOC,ic2.ASSOC1,ic2.ASSOC2,ic2.SPORT1,ic2.CONFNO1,ic2.SPORT2,ic2.CONFNO2,ic2.SPORT3,
    ic2.CONFNO3,ic2.SPORT4,ic2.CONFNO4,hd.LOCALE,hd.VETURL,ic2.VET1,ic2.VET2,ic2.VET3,ic2.VET4,ic2.VET5,ic2.VET9,efd.UGENTERN,vef.PctEnrWh,vef.PctEnrBK ,
    vef.PctEnrHS,vef.PctEnrAP,vef.PctEnrAN,vef.PctEnrUn,vef.PctEnrNr,vef.PCTENRW,ic2.DISAB,ic2.DISABPCT,f19_a.F1C061,vf.F1STSVPC,vf.F1STSVFT,vf.F2STSVPC,
    vf.F2STSVFT,hd.HBCU,hd.ACT,hd.CYACTIVE,hd.INSTNM,hd.CHFNM,hd.CHFTITLE, hd.GENTELE,hd.FAIDURL, hd.ADMINURL,hd.APPLURL,hd.DISAURL,hd.WEBADDR,hd.NPRICURL,
    ic2.CNTLAFFI,ic2.RELAFFIL,ic2.STUSRV1, ic2.STUSRV2, ic2.STUSRV4,ic2.STUSRV8,ic2.STUSRV9, ic2_m.MISSIONURL, ic2.DISTCRS,ic2.DISTNCED, ic2.DSTNCED1,
    ic2.DSTNCED3,ic2.DSTNUGC,ic2.DSTNUGP, ic2.DSTNUGN

FROM 
    HD2020 hd
LEFT JOIN 
    F1920_F1A f19_a ON hd.UNITID = f19_a.UNITID
LEFT JOIN 
    F1920_F2 f19_f2 ON f19_a.UNITID = f19_f2.UNITID
LEFT JOIN 
    IC2020 ic2 ON f19_f2.UNITID = ic2.UNITID
LEFT JOIN 
    IC2020_AY ic2_a ON ic2.UNITID = ic2_a.UNITID
LEFT JOIN 
    DRVIC2020 vic ON ic2_a.UNITID = vic.UNITID
LEFT JOIN 
    SFA1920_P2 ps ON ic2_a.UNITID = p2.UNITID
LEFT JOIN 
   IC2020_PY ic2_p ON p2.UNITID = ic2_p.UNITID
LEFT JOIN 
    ADM2020 adm ON ic2_p.UNITID = adm.UNITID
LEFT JOIN 
    EF2020D efd ON adm.UNITID = efd.UNITID
LEFT JOIN 
    DRVF2020 vf ON efd.UNITID = vf.UNITID
LEFT JOIN 
    DRVEF2020 vef ON vf.UNITID = vef.UNITID
LEFT JOIN 
    IC2020MISSION ic2_m ON vef.UNITID = ic2_m.UNITID;

""")

# Execute the query
result = conn.execute(sql_query)


def check_table(table_name):
    with conn:
        result = conn.execute(text(f"SELECT UNITID FROM {table_name}"))

        prev_value = None
        seen_values = {}
        x = 0
      
        for row in result:
            x +=1
            value = row[0]

            if value is None:
                print(f"null value in {table_name}", prev_value)

            elif value in seen_values:
                seen_values[value] += 1
                
            else:
                seen_values[value] = 1

            prev_value = value

        print(f'number of rows is {x}')
        print(f'number of duplicted variables is {len(seen_values)}')
      
        # Print the dictionary of duplicates after processing the table
        duplicates = {key: val for key, val in seen_values.items() if val > 1}
        if duplicates:
            print(f"Duplicates in {table_name}: {duplicates}")
        else:
            print(f"No duplicates found in {table_name}")

#function to fetch UNITID values for tables
def fetch_unitid_values(table_name):
    with engine.connect() as conn:
        result = conn.execute(text(f"SELECT UNITID FROM {table_name}"))
        return {row[0] for row in result}


#----------FUNCTION CALLING/ CODE TO CHECK DUPLICATE AND NULL VALUE UNITID'---------------
## Iterate over table names and check each one, excluding specified tables
# for table in table_names:
#     if table not in tables_not_used:
#         print(f"Checking table: {table}")
#         check_table(table)
#         print("Done checking", table)
# print(table_names)
#-------------------------------------------------------------------------------------------


#----------CODE TO COMPARE MISSING UNITID's FROM THE REFERENCE TABLE------------------------
## Fetch UNITID values from the reference table
# unitid_reference = fetch_unitid_values(reference_table)

# # Compare UNITID values for each table
# for table in other_tables:
#     unitid_other = fetch_unitid_values(table)

#     # Values in reference table but not in the other table
#     unique_in_reference = unitid_reference - unitid_other #don't know if I should use a minus or the difference () function
 
#     # Values in the other table but not in the reference table
#     unique_in_other = unitid_other - unitid_reference

#     # Print results
#     if unique_in_reference:
#         print(f"UNITID values in {reference_table} but not in {table}: {unique_in_reference}")
#     if unique_in_other:
#         print(f"UNITID values in {table} but not in {reference_table}: {unique_in_other}")

# print("Finished comparing UNITID values across tables.")
#-------------------------------------------------------------------------------------------  














# # Close the connection
# conn.close()
